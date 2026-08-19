# -*- coding: utf-8 -*-
"""Create a small Word contract used by the end-to-end upload smoke test."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def main():
    out = Path(__file__).resolve().parents[1] / "tmp" / "test-word-contract.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.right_margin = sec.bottom_margin = sec.left_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("WORD 合同上传测试"), 23, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(subtitle.add_run("电子印章服务端到端验证文档"), 13, False, "555555")

    for label, value in (("甲方", "测试甲方有限公司"), ("乙方", "测试乙方有限公司"), ("日期", "2026年8月19日")):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(label + "："), 11, True)
        set_font(p.add_run(value), 11)

    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(8)
    set_font(h.add_run("一、服务说明"), 16, True, "2E74B5")
    doc.add_paragraph("本文件用于验证系统能够上传 Word 合同、自动转换为 PDF、添加标准 4cm 电子印章并导出最终 PDF。")
    doc.add_paragraph("双方确认：测试结果仅用于软件功能验收，不构成真实商业合同。")

    sign = doc.add_paragraph()
    sign.paragraph_format.space_before = Pt(36)
    set_font(sign.add_run("甲方（盖章）：____________________"), 11, True)
    sign = doc.add_paragraph()
    sign.paragraph_format.space_before = Pt(24)
    set_font(sign.add_run("乙方（盖章）：____________________"), 11, True)

    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
